from docx import Document
from docx.table import Table
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.table import CT_Tbl, qn
from docx.shared import Inches, Pt, RGBColor
from bs4 import BeautifulSoup
import win32com.client
from win32com.client import constants as wd
import asyncio

alignment_map = {
   WD_ALIGN_PARAGRAPH.LEFT: "left",
   WD_ALIGN_PARAGRAPH.RIGHT: "right",
   WD_ALIGN_PARAGRAPH.CENTER: "center",
   WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
   WD_ALIGN_PARAGRAPH.DISTRIBUTE: "justify",
}
def get_style(paragraph):
    css_style = ''
    font_size = None
    if len(paragraph.runs)>0:
        font = paragraph.runs[0].font
        if font.name:
            css_style += f"font-family: {font.name}; "
        if font.size:
            css_style += f"font-size: {font.size.pt}pt; "
            font_size = font.size
        if font.bold:
            css_style += "font-weight: bold; "
        if font.italic:
            css_style += "font-style: italic; "
        if font.underline:
            css_style += "text-decoration: underline; "

    if paragraph.alignment and paragraph.alignment in alignment_map:
        css_style += f"text-align: {alignment_map[paragraph.alignment]};"
    else: css_style += "text-align: justify;"

    if font_size:
        if paragraph.paragraph_format.space_before:
            css_style += f"margin-top: {paragraph.paragraph_format.space_before/font_size}em;"
        if paragraph.paragraph_format.space_after:
            css_style += f"margin-bottom: {paragraph.paragraph_format.space_after/font_size}em;"
        if paragraph.paragraph_format.first_line_indent:
            css_style += f"text-indent: {max(paragraph.paragraph_format.first_line_indent/font_size,0)}em;"
    return css_style


def table_to_html(table, id):
    html = "<table border='1' cellspacing='0'>"
    css = ''
    cnt = 0
    cell_set = []
    for row in table.rows:
        html += "<tr>"
        for cell in row.cells:
            if cell in cell_set:
                html += '<td></td>'
                continue
            cell_set.append(cell)
            buf = ''
            for p in cell.paragraphs:
                buf += f"<p id=\"t{id}p{cnt}\">{p.text}</p>"
                css += f"#t{id}p{cnt}"+'{'+ get_style(p) + 'padding: 8px;' +'}\n'
                cnt += 1
            html += '<td>'+buf+'</td>'
        html += "</tr>"
    html += "</table>"
    return html, css


def paragraph_to_html(paragraph, id):
    text = paragraph.text
    if len(text)==0:
        return '',''
    html = f"<p id=\"p{id}\">{text}</p>"
    css  = f"#p{id}"+'{'+ get_style(paragraph) +'}\n'
    return html, css


def word_to_html(path):
    docx_file = Document(path)
    contents=''
    style_sheet=''
    pid, tip = 0, 0
    for element in docx_file.element.body:
        if isinstance(element, CT_Tbl):
            table = Table(element, docx_file)
            table_text, css = table_to_html(table,tip)
            contents += table_text
            style_sheet += css
            tip += 1
        if isinstance(element, CT_P):
            paragraph = Paragraph(element, docx_file)
            paragraph_text, css = paragraph_to_html(paragraph,pid)
            contents += paragraph_text
            style_sheet += css
            pid += 1
    return '<style>'+style_sheet+'</style>' + contents

##############################################

def add_revision(paragraph, text):
    paragraph.add_run(' ')
    run = paragraph.add_run(text)
    new_font = paragraph.runs[0].font
    run.font.name = new_font.name
    run.font.size = new_font.size
    run.font.underline = True
    run.font.color.rgb = RGBColor(176, 91, 219)


def add_comment(paragraph, comment, author):
    range = paragraph.Range
    Comment = range.Comments.Add(range)
    Comment.Range.Text = comment
    Comment.Author = author


async def html_to_word2(download_queue: asyncio.Queue, word_app):
    while True:
        html, path, author, event = await download_queue.get()
        print('收到下载命令')
    
        path = path.replace('C:','')
        soup = BeautifulSoup(html, 'lxml')
        docx_file = word_app.Documents.Open(path)
        
        id = 0
        for p in docx_file.paragraphs:
            content, mark = None, None
            while id < len(docx_file.paragraphs):
                try: 
                    content = soup.find('p', id=f'p{id}')
                    mark = content.find('mark')
                    mark.text
                    break
                except: id += 1
            
            if content and mark and content.text[:content.text.find(mark.text)-len('   ')] in p.Range.Text:
                add_comment(p, mark.text, author)
                id += 1
        
        for id, t in enumerate(docx_file.tables):
            cnt = 0
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.Range.paragraphs:
                        content = None
                        try: content = soup.find('p', id=f't{id}p{cnt}').find('mark')
                        except: pass
                        if content != None:
                            add_comment(p, content.text, author)
                        cnt += 1
                    
        docx_file.SaveAs(path+'.download')
        docx_file.Close()
        download_queue.task_done()
        event.set()


def html_to_word(html, path, suggestions):
    soup = BeautifulSoup(html, 'lxml')
    docx_file = Document(path)
    visit = {}

    for s in suggestions:
        if s.status == 1:
            visit[s.citation] = s.annotation

    for id, p in enumerate(docx_file.paragraphs):
        if f'p{id}' in visit:
            add_revision(p, visit[f'p{id}'])
    
    for id, t in enumerate(docx_file.tables):
        cnt = 0
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if f't{id}p{cnt}' in visit:
                        add_revision(p,'\n'+ visit[f't{id}p{cnt}'])
                        cell.width = Inches(1)
                    cnt += 1
                
    docx_file.save(path+'.download')


def written_to_word(content:str, path):
    doc = Document()
    doc.styles['Normal'].font.name = u'仿宋'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
    doc.styles['Normal'].font.size = Pt(14)
    parts = content.split('<br>')
    for part in parts:
        paragraph = doc.add_paragraph(part)
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.save(path+'.download')


if __name__=='__main__':
    pass